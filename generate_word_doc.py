#!/usr/bin/env python3
"""
Script to generate a Word document explaining the code flow and including all source code.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading_with_style(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.style.font.size = Pt(14 if level == 1 else 12)
    return heading

def add_code_block(doc, code, language="python"):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    para.style.font.name = 'Courier New'
    para.style.font.size = Pt(9)
    run = para.add_run(code)
    run.font.name = 'Courier New'
    return para

def create_word_document():
    """Create the comprehensive Word document"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Privacy-Preserving Smart Contracts in Cloud Computing', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Code Flow Documentation and Source Code')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Table of Contents placeholder
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. System Overview',
        '2. Architecture and Flow',
        '3. Application Structure',
        '4. Detailed Code Flow',
        '5. Source Code Files',
        '6. Key Components Explanation'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== SECTION 1: SYSTEM OVERVIEW ==========
    doc.add_heading('1. System Overview', 1)
    
    doc.add_paragraph(
        'This Privacy-Preserving Smart Contracts system is a Django-based web application '
        'that implements secure data access management using advanced cryptographic techniques. '
        'The system enables contract owners to create and manage data access policies, while '
        'requesters can request access to encrypted data through a secure validation process.'
    )
    
    doc.add_heading('1.1 Core Components', 2)
    components = [
        ('Contract Management', 'Enables owners to create contracts with access policies and encrypted file storage'),
        ('Request Processing', 'Handles data access requests with justification and approval workflow'),
        ('Secure Computation Layer', 'Implements ZKP (Zero-Knowledge Proofs), TEE (Trusted Execution Environment), and SMPC (Secure Multi-Party Computation)'),
        ('Oracle Service', 'Provides independent third-party verification and digital signing of approved requests'),
        ('Access Proxy', 'Verifies attestations and provides controlled data access with ephemeral decryption'),
        ('Encrypted Storage', 'Fernet-based symmetric encryption for data at rest'),
        ('Audit Logging', 'Comprehensive event tracking for compliance and forensics')
    ]
    
    for comp_name, comp_desc in components:
        p = doc.add_paragraph(comp_name, style='List Bullet')
        p.add_run(f': {comp_desc}')
    
    # ========== SECTION 2: ARCHITECTURE AND FLOW ==========
    doc.add_heading('2. Architecture and Flow', 1)
    
    doc.add_heading('2.1 System Flow Diagram', 2)
    doc.add_paragraph(
        'The system follows this workflow:'
    )
    
    flow_steps = [
        '1. User Registration/Login: Users register and authenticate through Django\'s authentication system',
        '2. Contract Creation: Contract owners create contracts with policies and upload encrypted files',
        '3. Contract Visibility: Contracts can be PUBLIC (visible to all) or PRIVATE (restricted to specific companies)',
        '4. Access Request: Data requesters browse public/private contracts and submit access requests with justification',
        '5. Owner Review: Contract owners review pending requests',
        '6. Secure Computation Validation: When approved, the system performs ZKP, TEE, and SMPC validations',
        '7. Oracle Attestation: Approved requests are signed by the Oracle service with cryptographic signatures',
        '8. Access Grant: Requesters can retrieve encrypted data through the Access Proxy after attestation verification',
        '9. Audit Logging: All events are logged for compliance and security monitoring'
    ]
    
    for step in flow_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('2.2 Data Flow', 2)
    doc.add_paragraph(
        'Data flows through the system as follows:'
    )
    doc.add_paragraph(
        '• File Upload → Encryption (Fernet) → Encrypted Storage → Contract Association',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Access Request → Secure Validation (ZKP/TEE/SMPC) → Oracle Attestation → Access Proxy → Decryption → User',
        style='List Bullet'
    )
    
    # ========== SECTION 3: APPLICATION STRUCTURE ==========
    doc.add_heading('3. Application Structure', 1)
    
    doc.add_paragraph('The Django project consists of the following apps:')
    
    apps_structure = """
privacy_smartcontracts/
├── contracts/          # Contract/policy management
│   ├── models.py       # Contract and ContractDocument models
│   ├── views.py        # Contract CRUD operations
│   ├── forms.py        # Contract creation/editing forms
│   └── urls.py         # URL routing
├── storage/            # Encrypted file storage
│   ├── models.py       # StoredObject model
│   ├── views.py        # Storage views
│   └── utils.py        # Encryption/decryption utilities
├── requests_app/       # Data access request handling
│   ├── models.py       # DataAccessRequest model
│   ├── views.py        # Request processing views
│   └── forms.py        # Request forms
├── secure_computation/ # ZKP Engine, TEE Gateway, SMPC Node
│   ├── models.py       # SecureComputationValidation model
│   └── views.py        # Validation dashboard
├── oracle/             # Mock oracle to sign requests
│   ├── models.py       # Attestation model
│   └── views.py        # Signing operations
├── access_proxy/       # Retrieves file after verifying attestation
│   └── views.py        # Attestation verification and file retrieval
├── audit/              # Logs all events
│   ├── models.py       # AuditEvent model
│   ├── views.py        # Audit log views
│   └── utils.py        # Logging utilities
├── users/              # User accounts and roles
│   ├── models.py       # UserProfile model
│   ├── views.py        # Registration views
│   └── forms.py        # User forms
└── privacy_smartcontracts/  # Main Django project
    ├── settings.py     # Django settings
    └── urls.py         # Main URL configuration
"""
    
    add_code_block(doc, apps_structure)
    
    # ========== SECTION 4: DETAILED CODE FLOW ==========
    doc.add_heading('4. Detailed Code Flow', 1)
    
    doc.add_heading('4.1 User Authentication Flow', 2)
    doc.add_paragraph(
        'Users register through the users app, which creates a User account and optionally a UserProfile '
        'with company information. Login uses Django\'s built-in authentication system.'
    )
    
    doc.add_heading('4.2 Contract Creation Flow', 2)
    doc.add_paragraph(
        '1. Owner navigates to contract creation page',
        style='List Number'
    )
    doc.add_paragraph(
        '2. Owner fills ContractForm with title, description, visibility, and policy details',
        style='List Number'
    )
    doc.add_paragraph(
        '3. Contract is saved with owner_accepted=True',
        style='List Number'
    )
    doc.add_paragraph(
        '4. Owner can upload documents which are encrypted using Fernet and stored as StoredObject',
        style='List Number'
    )
    doc.add_paragraph(
        '5. Contract status updates based on owner and second_party acceptance',
        style='List Number'
    )
    
    doc.add_heading('4.3 Access Request Flow', 2)
    doc.add_paragraph(
        '1. Requester browses public/private contracts',
        style='List Number'
    )
    doc.add_paragraph(
        '2. Requester submits DataAccessRequest with justification',
        style='List Number'
    )
    doc.add_paragraph(
        '3. Request status is PENDING',
        style='List Number'
    )
    doc.add_paragraph(
        '4. Owner reviews request in owner_requests view',
        style='List Number'
    )
    doc.add_paragraph(
        '5. Owner approves/denies request',
        style='List Number'
    )
    
    doc.add_heading('4.4 Secure Computation Validation Flow', 2)
    doc.add_paragraph(
        'When a request is approved, the system automatically performs secure computation validation:'
    )
    doc.add_paragraph(
        '1. SecureComputationValidation object is created/retrieved',
        style='List Number'
    )
    doc.add_paragraph(
        '2. ZKP verification: Simplified zero-knowledge proof is generated',
        style='List Number'
    )
    doc.add_paragraph(
        '3. TEE validation: TEEGateway performs secure computation and generates remote attestation with ECDSA signature',
        style='List Number'
    )
    doc.add_paragraph(
        '4. SMPC verification: Secure multi-party computation result is generated',
        style='List Number'
    )
    doc.add_paragraph(
        '5. Overall verification status is set based on all three validations',
        style='List Number'
    )
    
    doc.add_heading('4.5 Oracle Attestation Flow', 2)
    doc.add_paragraph(
        '1. After secure validation passes, Oracle creates an Attestation',
        style='List Number'
    )
    doc.add_paragraph(
        '2. Payload is created with request details',
        style='List Number'
    )
    doc.add_paragraph(
        '3. Payload is signed using Ed25519 private key',
        style='List Number'
    )
    doc.add_paragraph(
        '4. Signature is base64 encoded and stored with attestation',
        style='List Number'
    )
    
    doc.add_heading('4.6 Data Access Flow', 2)
    doc.add_paragraph(
        '1. Requester clicks retrieve link for approved request',
        style='List Number'
    )
    doc.add_paragraph(
        '2. Access Proxy verifies attestation signature using Oracle public key',
        style='List Number'
    )
    doc.add_paragraph(
        '3. If valid, encrypted file is retrieved from StoredObject',
        style='List Number'
    )
    doc.add_paragraph(
        '4. File is decrypted using Fernet key',
        style='List Number'
    )
    doc.add_paragraph(
        '5. Decrypted file is sent to requester',
        style='List Number'
    )
    doc.add_paragraph(
        '6. Access event is logged in audit system',
        style='List Number'
    )
    
    # ========== SECTION 5: SOURCE CODE FILES ==========
    doc.add_heading('5. Source Code Files', 1)
    
    # Read and include all source files
    base_path = 'privacy_smartcontracts'
    
    source_files = [
        ('Main Settings', 'privacy_smartcontracts/privacy_smartcontracts/settings.py'),
        ('Main URLs', 'privacy_smartcontracts/privacy_smartcontracts/urls.py'),
        ('Contract Models', 'privacy_smartcontracts/contracts/models.py'),
        ('Contract Views', 'privacy_smartcontracts/contracts/views.py'),
        ('Contract Forms', 'privacy_smartcontracts/contracts/forms.py'),
        ('Contract URLs', 'privacy_smartcontracts/contracts/urls.py'),
        ('Storage Models', 'privacy_smartcontracts/storage/models.py'),
        ('Storage Utils', 'privacy_smartcontracts/storage/utils.py'),
        ('Request Models', 'privacy_smartcontracts/requests_app/models.py'),
        ('Request Views', 'privacy_smartcontracts/requests_app/views.py'),
        ('Request Forms', 'privacy_smartcontracts/requests_app/forms.py'),
        ('Secure Computation Models', 'privacy_smartcontracts/secure_computation/models.py'),
        ('Secure Computation Views', 'privacy_smartcontracts/secure_computation/views.py'),
        ('Oracle Models', 'privacy_smartcontracts/oracle/models.py'),
        ('Oracle Views', 'privacy_smartcontracts/oracle/views.py'),
        ('Access Proxy Views', 'privacy_smartcontracts/access_proxy/views.py'),
        ('Audit Models', 'privacy_smartcontracts/audit/models.py'),
        ('Audit Views', 'privacy_smartcontracts/audit/views.py'),
        ('Audit Utils', 'privacy_smartcontracts/audit/utils.py'),
        ('User Models', 'privacy_smartcontracts/users/models.py'),
        ('User Views', 'privacy_smartcontracts/users/views.py'),
        ('User URLs', 'privacy_smartcontracts/users/urls.py'),
    ]
    
    for title, file_path in source_files:
        doc.add_heading(f'5.{source_files.index((title, file_path)) + 1} {title}', 2)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                code_content = f.read()
            add_code_block(doc, code_content)
        except FileNotFoundError:
            doc.add_paragraph(f'File not found: {file_path}', style='Intense Quote')
        except Exception as e:
            doc.add_paragraph(f'Error reading file: {str(e)}', style='Intense Quote')
        
        doc.add_paragraph()  # Spacing between files
    
    # ========== SECTION 6: KEY COMPONENTS EXPLANATION ==========
    doc.add_heading('6. Key Components Explanation', 1)
    
    doc.add_heading('6.1 Encryption System', 2)
    doc.add_paragraph(
        'The system uses Fernet (symmetric encryption) from the cryptography library. '
        'Files are encrypted when uploaded and stored as .enc files. The FERNET_KEY is stored '
        'in environment variables for security.'
    )
    
    doc.add_heading('6.2 TEE Implementation', 2)
    doc.add_paragraph(
        'The TEEGateway class simulates a Trusted Execution Environment using real cryptographic primitives:'
    )
    doc.add_paragraph(
        '• Uses ECDSA (SECP256R1) for attestation signing',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Generates enclave IDs and measurements (simulating TPM/SGX)',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Creates remote attestation quotes with cryptographic signatures',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Verifies attestations using public key cryptography',
        style='List Bullet'
    )
    
    doc.add_heading('6.3 Oracle Attestation', 2)
    doc.add_paragraph(
        'The Oracle service uses Ed25519 digital signatures for attestations. The private key '
        'is stored in environment variables, and the public key is used for verification.'
    )
    
    doc.add_heading('6.4 Audit System', 2)
    doc.add_paragraph(
        'All system events are logged through the audit.utils.log_event() function. Events include:'
    )
    doc.add_paragraph(
        '• Contract creation and updates',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Access request creation and processing',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Secure computation validations',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Attestation issuance',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Data access events',
        style='List Bullet'
    )
    
    doc.add_heading('6.5 Security Features', 2)
    doc.add_paragraph(
        'The system implements multiple security layers:'
    )
    doc.add_paragraph(
        '• File encryption at rest using Fernet',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Secure computation validation (ZKP/TEE/SMPC)',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Cryptographic attestations with digital signatures',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Access control through contract visibility and policies',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Comprehensive audit logging',
        style='List Bullet'
    )
    
    # Save document
    output_file = 'Code_Flow_Documentation.docx'
    doc.save(output_file)
    print(f'Document created successfully: {output_file}')
    return output_file

if __name__ == '__main__':
    try:
        create_word_document()
    except ImportError:
        print("Error: python-docx library is required.")
        print("Please install it using: pip install python-docx")
    except Exception as e:
        print(f"Error creating document: {e}")
        import traceback
        traceback.print_exc()

